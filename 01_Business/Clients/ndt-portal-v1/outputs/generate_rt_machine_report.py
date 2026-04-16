"""
Generate RT Machine Technical Specifications — Client Report (.docx)
Run: python outputs/generate_rt_machine_report.py
Output: outputs/RT_Machine_Technical_Specifications_Report.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

OUT_PATH = os.path.join(os.path.dirname(__file__), "RT_Machine_Technical_Specifications_Report.docx")

# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY    = RGBColor(0x1A, 0x2E, 0x4A)   # headers
BLUE    = RGBColor(0x1F, 0x5C, 0x99)   # sub-headers / accents
AMBER   = RGBColor(0xD9, 0x7B, 0x06)   # supplemented fields callout
DARK    = RGBColor(0x1F, 0x27, 0x37)   # body text
MUTED   = RGBColor(0x6B, 0x72, 0x80)   # captions / notes
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE = RGBColor(0xEB, 0xF3, 0xFB)  # table header bg
LIGHT_GRAY = RGBColor(0xF8, 0xF9, 0xFA)  # alternating rows

# ── Helpers ────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)

def set_cell_border(cell, top=False, bottom=False, color='C8D3DF'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, active in [('top', top), ('bottom', bottom)]:
        if active:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:color'), color)
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18) if level == 1 else Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = NAVY
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = BLUE
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
    p.paragraph_format.keep_with_next = True
    return p

def add_body(doc, text, italic=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    run.italic = italic
    return p

def add_callout(doc, text, style='info'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = AMBER if style == 'warning' else BLUE
    run.italic = True
    return p

def add_kv_table(doc, rows, asterisk_keys=None):
    """Add a 2-col key-value table. asterisk_keys marks * supplemented fields."""
    asterisk_keys = asterisk_keys or set()
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Cm(7)
    tbl.columns[1].width = Cm(9.5)
    for i, (k, v) in enumerate(rows):
        row = tbl.rows[i]
        row.height = Cm(0.55)
        # key cell
        c0 = row.cells[0]
        set_cell_bg(c0, LIGHT_GRAY if i % 2 == 0 else WHITE)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(1)
        p0.paragraph_format.space_after  = Pt(1)
        kr = p0.add_run(k)
        kr.font.size = Pt(9)
        kr.font.color.rgb = MUTED
        if k in asterisk_keys:
            sup = p0.add_run(' *')
            sup.font.size = Pt(7)
            sup.font.color.rgb = AMBER
        # value cell
        c1 = row.cells[1]
        set_cell_bg(c1, LIGHT_GRAY if i % 2 == 0 else WHITE)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(1)
        p1.paragraph_format.space_after  = Pt(1)
        vr = p1.add_run(str(v))
        vr.font.size = Pt(9)
        vr.font.color.rgb = DARK
        vr.bold = k in asterisk_keys
    doc.add_paragraph()
    return tbl

def add_comparison_table(doc, headers, data_rows):
    col_count = len(headers)
    tbl = doc.add_table(rows=1 + len(data_rows), cols=col_count)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = tbl.rows[0]
    hrow.height = Cm(0.7)
    for j, h in enumerate(headers):
        c = hrow.cells[j]
        set_cell_bg(c, NAVY)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(h)
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = WHITE

    for i, row_data in enumerate(data_rows):
        row = tbl.rows[i + 1]
        row.height = Cm(0.55)
        for j, val in enumerate(row_data):
            c = row.cells[j]
            set_cell_bg(c, LIGHT_GRAY if i % 2 == 0 else WHITE)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.color.rgb = DARK
            if j == 0:
                r.bold = True

    doc.add_paragraph()

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), 'C8D3DF')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Machine data ───────────────────────────────────────────────────────────────
MACHINES = [
    {
        "title":      "Unit 1 — Varian NDI 320/26  (Walk-In Vault)",
        "machine_id": "RT_VARIAN_320_26",
        "make_model": "Varian / NDI 320/26",
        "chamber_type": "Walk-In Vault",
        "supplemented_note": (
            "Fields marked * were supplemented from the Varex NDI-320-26 Product Data Sheet "
            "(vareximaging.com) and xrayllc.com distributor datasheet, as the customer's "
            "specification response did not include these tube-level parameters."
        ),
        "source_rows": [
            ("xray_source", [
                ("Type",                      "NDT — Industrial X-ray"),
                ("Maximum voltage",            "320 kV"),
                ("Operating range",            "40 – 300 kV  (customer-provided)"),
                ("Max mA at rated voltage",    "10 mA  (customer-provided)"),
                ("Focal spot class",           "Standard industrial"),
                ("Focal spot size (small) *",  "1.5 mm  (IEC 336 / EN 12543)"),
                ("Focal spot size (large) *",  "4.0 mm  (IEC 336 / EN 12543)"),
                ("Beam cone angle *",          "40 degrees"),
                ("Target / anode angle *",     "20 degrees"),
                ("Target material *",          "Tungsten — stationary anode, metal-ceramic"),
                ("Inherent filtration *",      "4.0 mm Be  (beryllium window)"),
                ("Max continuous power (small focal) *", "1,500 W  (@ 14 L/min oil flow)"),
                ("Max continuous power (large focal) *", "4,200 W  (@ 14 L/min oil flow)"),
                ("Cooling *",                  "Oil cooled — 14 L/min minimum flow"),
                ("Modality",                   "Film RT  (film_rt)"),
            ]),
            ("inspection_envelope", [
                ("Chamber shape",              "Rectangular — Walk-In Vault"),
                ("Max part diameter",          "914.4 mm  (36 inches)"),
                ("Max part height",            "1,526 mm  (60 inches)"),
                ("Max part weight",            "22.7 kg  (50 lbs)"),
                ("Usable clearance",           "36 inches — part + fixturing"),
                ("FFD (Focus-to-Film Distance)","64 inches  (machine log reference)"),
            ]),
            ("manipulation", [
                ("Available axes",             "Rotate, Vertical, Horizontal, Tilt"),
                ("Tilt available",             "Yes"),
                ("Rotation range",             "0 – 90 degrees"),
            ]),
            ("detector_support", [
                ("Film supported",             "Yes"),
                ("Digital detector",           "No"),
                ("Typical film classes",       "DR, D2, D3, D4, D5, D7, IX50, D4 ETE PB"),
                ("IQI support",                "IQI Wire, IQI Hole"),
            ]),
            ("planning_rules", [
                ("Best for",                   "Large parts, long pipe spools, parts requiring tilt, multi-axis manipulation"),
                ("Not ideal for",              "Gamma source jobs (Ir-192 / Co-60 not applicable to this cabinet)"),
            ]),
        ],
        "asterisk_keys": {
            "Focal spot size (small) *", "Focal spot size (large) *",
            "Beam cone angle *", "Target / anode angle *",
            "Target material *", "Inherent filtration *",
            "Max continuous power (small focal) *",
            "Max continuous power (large focal) *",
            "Cooling *",
        },
        "refs": [
            ("Varex NDI-320-26 Product Data Sheet (official)",
             "https://www.vareximaging.com/wp-content/uploads/2022/01/NDI-320-26_PDS_133112-000.pdf"),
            ("xrayllc.com — NDI 320/26 Distributor Datasheet",
             "https://xrayllc.com/ndi320-26.pdf"),
        ],
    },
    {
        "title":      "Unit 2 — Comet MXR225/22  (Cabinet)",
        "machine_id": "RT_COMET_MXR225_22",
        "make_model": "Comet MXR225/22",
        "chamber_type": "Cabinet",
        "supplemented_note": (
            "Fields marked * were supplemented from the Comet Group official product page "
            "(xray.comet.tech) and a verified distributor technical listing."
        ),
        "source_rows": [
            ("xray_source", [
                ("Type",                      "NDT — Industrial X-ray"),
                ("Maximum voltage",            "225 kV"),
                ("Operating range",            "40 – 220 kV  (customer-provided)"),
                ("Max mA at rated voltage",    "10 mA  (customer-provided)"),
                ("Focal spot class",           "Standard industrial"),
                ("Focal spot size (small) *",  "1.0 mm  (EN 12543)"),
                ("Focal spot size (large) *",  "5.5 mm  (EN 12543)"),
                ("Beam cone angle *",          "40 degrees"),
                ("Target / anode angle *",     "20 degrees"),
                ("Target material *",          "Tungsten — stationary anode, metal-ceramic, unipolar"),
                ("Inherent filtration *",      "0.8 ± 0.1 mm Be  (beryllium window)"),
                ("Max continuous power (small focal) *", "640 W"),
                ("Max continuous power (large focal) *", "3,000 W"),
                ("Cooling *",                  "Water cooled — 4 L/min minimum flow"),
                ("Modality",                   "Film RT  (film_rt)"),
            ]),
            ("inspection_envelope", [
                ("Chamber shape",              "Rectangular — Cabinet"),
                ("Max part diameter",          "304.8 mm  (12 inches)"),
                ("Max part height",            "304.8 mm  (12 inches)"),
                ("Max part weight",            "22.68 kg"),
                ("Usable clearance",           "12 inches"),
                ("FFD (Focus-to-Film Distance)","44 inches  (machine log reference)"),
            ]),
            ("manipulation", [
                ("Available axes",             "Vertical only"),
                ("Tilt available",             "No"),
                ("Rotation",                   "N/A"),
            ]),
            ("detector_support", [
                ("Film supported",             "Yes"),
                ("Digital detector",           "No"),
                ("Typical film classes",       "DR, D2, D3, D4, D5, D7, IX50, D4 ETE PB"),
                ("IQI support",                "IQI Wire, IQI Hole"),
            ]),
            ("planning_rules", [
                ("Best for",                   "Small parts, compact geometry, lower-energy jobs"),
                ("Not ideal for",              "Gamma source jobs; parts exceeding 304.8 mm envelope"),
            ]),
        ],
        "asterisk_keys": {
            "Focal spot size (small) *", "Focal spot size (large) *",
            "Beam cone angle *", "Target / anode angle *",
            "Target material *", "Inherent filtration *",
            "Max continuous power (small focal) *",
            "Max continuous power (large focal) *",
            "Cooling *",
        },
        "refs": [
            ("Comet MXR-225/22 Official Product Page",
             "https://xray.comet.tech/en/products/mxr-225-22"),
            ("Comet MXR-225/22 Technical Datasheet — Distributor (made-in-china.com)",
             "https://zxt-xray.en.made-in-china.com/product/CdIAbmflHxco/China-Comet-X-ray-Tube-for-X-ray-Machine-225kv-Mxr-225-22-.html"),
        ],
    },
    {
        "title":      "Unit 3 — Comet MXR320/26  (Cabinet)",
        "machine_id": "RT_COMET_MXR320_26",
        "make_model": "Comet MXR320/26",
        "chamber_type": "Cabinet",
        "supplemented_note": (
            "Fields marked * were supplemented from the Comet Group official product page "
            "(xray.comet.tech) and the Willick Engineering distributor datasheet."
        ),
        "source_rows": [
            ("xray_source", [
                ("Type",                      "NDT — Industrial X-ray"),
                ("Maximum voltage",            "320 kV"),
                ("Operating range",            "40 – 300 kV  (customer-provided)"),
                ("Max mA at rated voltage",    "10 mA  (customer-provided)"),
                ("Focal spot class",           "Standard industrial"),
                ("Focal spot size (small) *",  "3.0 mm  (EN 12543)"),
                ("Focal spot size (large) *",  "5.5 mm  (EN 12543)"),
                ("Beam cone angle *",          "40 degrees"),
                ("Target / anode angle *",     "20 degrees"),
                ("Target material *",          "Tungsten — stationary anode, metal-ceramic, bipolar"),
                ("Inherent filtration *",      "3.0 mm Be  (beryllium window)"),
                ("Max continuous power (small focal) *", "1,500 W"),
                ("Max continuous power (large focal) *", "4,200 W"),
                ("Cooling *",                  "Oil cooled — 14 L/min minimum, max inlet temp 50 °C"),
                ("Modality",                   "Film RT  (film_rt)"),
            ]),
            ("inspection_envelope", [
                ("Chamber shape",              "Rectangular — Cabinet"),
                ("Max part diameter",          "457.2 mm  (18 inches)"),
                ("Max part height",            "304.8 mm  (12 inches)"),
                ("Max part weight",            "22.68 kg"),
                ("Usable clearance",           "12 inches"),
                ("FFD (Focus-to-Film Distance)","44 inches  (machine log reference)"),
            ]),
            ("manipulation", [
                ("Available axes",             "Vertical only"),
                ("Tilt available",             "No"),
                ("Rotation",                   "N/A"),
            ]),
            ("detector_support", [
                ("Film supported",             "Yes"),
                ("Digital detector",           "No"),
                ("Typical film classes",       "DR, D2, D3, D4, D5, D7, IX50, D4 ETE PB"),
                ("IQI support",                "IQI Wire, IQI Hole"),
            ]),
            ("planning_rules", [
                ("Best for",                   "Medium-diameter parts, thick sections requiring 320 kV penetration"),
                ("Not ideal for",              "Gamma source jobs; parts taller than 304.8 mm"),
            ]),
        ],
        "asterisk_keys": {
            "Focal spot size (small) *", "Focal spot size (large) *",
            "Beam cone angle *", "Target / anode angle *",
            "Target material *", "Inherent filtration *",
            "Max continuous power (small focal) *",
            "Max continuous power (large focal) *",
            "Cooling *",
        },
        "refs": [
            ("Comet MXR-320/26 Official Product Page",
             "https://xray.comet.tech/en/products/mxr-320-26"),
            ("Comet MXR-320/26 Datasheet via Willick Engineering",
             "https://www.willick.com/wp-content/uploads/2021/05/comet-mxr-320_26_single_sheet_en_v6.pdf"),
        ],
    },
    {
        "title":      "Unit 4 — Varex NDI 320/26  (Cabinet)",
        "machine_id": "RT_VAREX_NDI320_26",
        "make_model": "Varex NDI 320/26",
        "chamber_type": "Cabinet",
        "supplemented_note": (
            "Fields marked * were supplemented from the Varex NDI-320-26 Product Data Sheet. "
            "Note: Varex Imaging acquired Varian Medical Systems' Imaging Components division in 2017. "
            "The NDI 320/26 is the same X-ray tube product now manufactured and sold under the Varex brand. "
            "All tube-level specifications are identical to Unit 1."
        ),
        "source_rows": [
            ("xray_source", [
                ("Type",                      "NDT — Industrial X-ray"),
                ("Maximum voltage",            "320 kV"),
                ("Operating range",            "40 – 300 kV  (customer-provided)"),
                ("Max mA at rated voltage",    "10 mA  (customer-provided)"),
                ("Focal spot class",           "Standard industrial"),
                ("Focal spot size (small) *",  "1.5 mm  (IEC 336 / EN 12543)"),
                ("Focal spot size (large) *",  "4.0 mm  (IEC 336 / EN 12543)"),
                ("Beam cone angle *",          "40 degrees"),
                ("Target / anode angle *",     "20 degrees"),
                ("Target material *",          "Tungsten — stationary anode, metal-ceramic"),
                ("Inherent filtration *",      "4.0 mm Be  (beryllium window)"),
                ("Max continuous power (small focal) *", "1,500 W  (@ 14 L/min oil flow)"),
                ("Max continuous power (large focal) *", "4,200 W  (@ 14 L/min oil flow)"),
                ("Cooling *",                  "Oil cooled — 14 L/min minimum flow"),
                ("Modality",                   "Film RT  (film_rt)"),
            ]),
            ("inspection_envelope", [
                ("Chamber shape",              "Rectangular — Cabinet"),
                ("Max part diameter",          "457.2 mm  (18 inches)"),
                ("Max part height",            "304.8 mm  (12 inches)"),
                ("Max part weight",            "22.68 kg"),
                ("Usable clearance",           "12 inches"),
                ("FFD (Focus-to-Film Distance)","48 inches  (machine log reference)"),
            ]),
            ("manipulation", [
                ("Available axes",             "Vertical only"),
                ("Tilt available",             "No"),
                ("Rotation",                   "N/A"),
            ]),
            ("detector_support", [
                ("Film supported",             "Yes"),
                ("Digital detector",           "No"),
                ("Typical film classes",       "DR, D2, D3, D4, D5, D7, IX50, D4 ETE PB"),
                ("IQI support",                "IQI Wire, IQI Hole"),
            ]),
            ("planning_rules", [
                ("Best for",                   "Medium-diameter parts, thick sections requiring 320 kV penetration"),
                ("Not ideal for",              "Gamma source jobs; parts taller than 304.8 mm"),
            ]),
        ],
        "asterisk_keys": {
            "Focal spot size (small) *", "Focal spot size (large) *",
            "Beam cone angle *", "Target / anode angle *",
            "Target material *", "Inherent filtration *",
            "Max continuous power (small focal) *",
            "Max continuous power (large focal) *",
            "Cooling *",
        },
        "refs": [
            ("Varex NDI-320-26 Product Data Sheet (official)",
             "https://www.vareximaging.com/wp-content/uploads/2022/01/NDI-320-26_PDS_133112-000.pdf"),
            ("Varex Imaging Corporate — former Varian NDT division",
             "https://www.vareximaging.com"),
        ],
    },
]

SECTION_TITLES = {
    "xray_source":          "1. X-Ray Source",
    "inspection_envelope":  "2. Inspection Envelope",
    "manipulation":         "3. Manipulation",
    "detector_support":     "4. Detector Support",
    "planning_rules":       "5. Planning Rules & Machine Suitability",
}


# ── Build document ─────────────────────────────────────────────────────────────
def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    # ── Cover block ────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run("RT MACHINE TECHNICAL SPECIFICATIONS")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run("Research & Configuration Report")
    r2.font.size = Pt(14)
    r2.font.color.rgb = BLUE

    p3 = doc.add_paragraph()
    r3 = p3.add_run(
        f"Prepared by: Onnex AI Agency  ·  NDT Portal v1 Project  ·  "
        f"{datetime.date.today().strftime('%B %d, %Y')}"
    )
    r3.font.size = Pt(9)
    r3.font.color.rgb = MUTED

    add_horizontal_rule(doc)

    # ── Executive Summary ──────────────────────────────────────────────────────
    add_heading(doc, "Executive Summary", 2)
    add_body(doc,
        "This report documents the technical specifications for four (4) industrial "
        "radiographic testing (RT) X-ray machines operated at the client's facility. "
        "Specifications were gathered from two sources:"
    )
    add_body(doc, "1.  Customer-provided data — submitted via machine specification request form.", indent=True)
    add_body(doc,
        "2.  Manufacturer datasheet research — supplemented from official Varex Imaging and "
        "Comet Group product documentation, where customer responses did not include "
        "tube-level parameters required for automated technique planning.",
        indent=True
    )
    add_body(doc,
        "Fields supplemented from external research are marked with an asterisk (*) throughout "
        "this document. All supplemented values are sourced from publicly available manufacturer "
        "datasheets and are referenced in the Sources section of each machine profile."
    )
    add_body(doc,
        "These specifications have been loaded into the NDT Portal v1 machine catalog and are "
        "used by the AI-assisted RT technique planner to select the appropriate machine, "
        "calculate geometric unsharpness (Ug), and generate technique cards per ASME Section V / "
        "ASTM E1742 requirements.",
        italic=True
    )

    add_horizontal_rule(doc)

    # ── Machine profiles ───────────────────────────────────────────────────────
    add_heading(doc, "Machine Profiles", 1)

    for m in MACHINES:
        doc.add_page_break()
        add_heading(doc, m["title"], 2)

        add_callout(doc,
            f"Machine ID: {m['machine_id']}  ·  Make/Model: {m['make_model']}  ·  "
            f"Chamber Type: {m['chamber_type']}",
            style='info'
        )
        add_callout(doc, f"★  {m['supplemented_note']}", style='warning')

        for section_key, rows in m["source_rows"]:
            add_heading(doc, SECTION_TITLES[section_key], 3)
            add_kv_table(doc, rows, m["asterisk_keys"])

        # Sources for this machine
        add_heading(doc, "Sources", 3)
        for src_name, url in m["refs"]:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            r_name = p.add_run(src_name + "  ")
            r_name.font.size = Pt(9)
            r_name.font.color.rgb = DARK
            r_url = p.add_run(url)
            r_url.font.size = Pt(8)
            r_url.font.color.rgb = BLUE

    # ── Comparative Analysis ───────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "Comparative Analysis", 1)
    add_body(doc,
        "The table below summarises the key differentiating parameters across all four machines. "
        "Parameters marked * are supplemented from manufacturer datasheets."
    )

    headers = ["Parameter", "Unit 1\nVarian NDI 320/26\n(Walk-In)", "Unit 2\nComet MXR225/22\n(Cabinet)", "Unit 3\nComet MXR320/26\n(Cabinet)", "Unit 4\nVarex NDI 320/26\n(Cabinet)"]
    data_rows = [
        ("Max kV",                     "320 kV",     "225 kV",     "320 kV",     "320 kV"),
        ("Max mA",                      "10 mA",      "10 mA",      "10 mA",      "10 mA"),
        ("FFD (machine log)",           "64 inches",  "44 inches",  "44 inches",  "48 inches"),
        ("Chamber type",               "Walk-In",    "Cabinet",    "Cabinet",    "Cabinet"),
        ("Max part diameter",          "914 mm",     "305 mm",     "457 mm",     "457 mm"),
        ("Max part height",            "1,526 mm",   "305 mm",     "305 mm",     "305 mm"),
        ("Axes",                       "4 (full)",   "Vertical",   "Vertical",   "Vertical"),
        ("Tilt",                       "Yes",        "No",         "No",         "No"),
        ("Focal spot — small *",       "1.5 mm",     "1.0 mm",     "3.0 mm",     "1.5 mm"),
        ("Focal spot — large *",       "4.0 mm",     "5.5 mm",     "5.5 mm",     "4.0 mm"),
        ("Beam cone angle *",          "40°",        "40°",        "40°",        "40°"),
        ("Target angle *",             "20°",        "20°",        "20°",        "20°"),
        ("Inherent filtration *",      "4.0 mm Be",  "0.8 mm Be",  "3.0 mm Be",  "4.0 mm Be"),
        ("Power — small focal *",      "1,500 W",    "640 W",      "1,500 W",    "1,500 W"),
        ("Power — large focal *",      "4,200 W",    "3,000 W",    "4,200 W",    "4,200 W"),
        ("Cooling *",                  "Oil",        "Water",      "Oil",        "Oil"),
        ("Film only",                  "Yes",        "Yes",        "Yes",        "Yes"),
    ]
    add_comparison_table(doc, headers, data_rows)

    # ── Planning Implications ──────────────────────────────────────────────────
    add_heading(doc, "Technical Planning Implications", 1)

    add_heading(doc, "Geometric Unsharpness (Ug)", 2)
    add_body(doc,
        "Ug is calculated as:  Ug = F × (OFD / SFD)  where F = focal spot size in mm, "
        "OFD = object-to-film distance, SFD = source-to-film distance."
    )
    add_body(doc,
        "The Comet MXR225/22 offers the smallest focal spot (1.0 mm small / 5.5 mm large). "
        "At small focal spot with equal geometry, it achieves the lowest Ug of all four machines — "
        "best for fine defect detection in thin-walled parts. "
        "The Comet MXR320/26 has the largest small focal spot (3.0 mm), producing the highest Ug "
        "at equivalent geometry — suitable for thick sections where penetration is the priority."
    )

    add_heading(doc, "Beam Coverage at Given FFD", 2)
    add_body(doc,
        "All four tubes share a 40° cone angle. Coverage diameter = 2 × SFD × tan(20°). "
        "At the machine-logged FFDs:"
    )
    add_body(doc, "Unit 1 (Walk-In) — FFD 64\":  coverage diameter ≈ 46.6 inches", indent=True)
    add_body(doc, "Unit 4 (Varex Cabinet) — FFD 48\":  coverage diameter ≈ 34.9 inches", indent=True)
    add_body(doc, "Units 2 & 3 (Comet Cabinets) — FFD 44\":  coverage diameter ≈ 32.0 inches", indent=True)

    add_heading(doc, "Beam Hardening & Filtration", 2)
    add_body(doc,
        "Inherent filtration varies significantly across machines. The Comet MXR225/22 has only "
        "0.8 mm Be filtration, producing a softer beam. For thick steel sections this may require "
        "additional external filtration (e.g. Cu or Al filters) to reduce scatter and beam hardening "
        "artefacts. The Varian/Varex NDI 320/26 tubes (4.0 mm Be) produce the most pre-hardened beam "
        "— less scatter, better contrast in dense materials."
    )

    add_heading(doc, "Power-Limited Exposure at Small Focal Spot", 2)
    add_body(doc,
        "At small focal spot, maximum continuous power limits effective mA at high kV. "
        "At 225 kV, the Comet MXR225/22 is limited to 640 W → max ~2.8 mA. "
        "At 320 kV, Varian/Varex and Comet MXR320/26 reach 1,500 W → max ~4.7 mA. "
        "Longer exposure times are required at small focal spot — a factor in technique card generation."
    )

    # ── Known Data Gaps ────────────────────────────────────────────────────────
    add_horizontal_rule(doc)
    add_heading(doc, "Known Data Gaps", 2)
    add_body(doc,
        "The following parameters were not available in public manufacturer documentation at "
        "time of research. Obtaining these from the client or manufacturer would improve "
        "technique planning accuracy:"
    )
    gaps = [
        "Radiation output curves (mGy/min or R/min at 1 m at rated kV/mA) — required for exposure time calculation",
        "mA minimum range and step resolution — generator specifications, not tube specs",
        "Duty cycle / maximum single-exposure time — operating manuals required",
        "Physical cabinet dimensions (L × W × H mm) — needed for access planning in confined geometries",
        "Generator model numbers for all four machines — may provide additional technique constraints",
    ]
    for g in gaps:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(g)
        r.font.size = Pt(9)
        r.font.color.rgb = DARK

    # ── Footer note ────────────────────────────────────────────────────────────
    add_horizontal_rule(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(
        "This document was prepared by Onnex AI Agency as part of the NDT Portal v1 engagement. "
        "Supplemented specifications are sourced from publicly available manufacturer datasheets "
        "and are provided for configuration purposes only. Customer-provided data takes precedence "
        "in all cases. If any specification conflicts with the client's calibrated equipment records, "
        "the calibrated records should be used."
    )
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED
    r.italic = True

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == '__main__':
    build()
