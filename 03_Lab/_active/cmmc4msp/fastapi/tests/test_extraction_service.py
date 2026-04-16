"""Unit tests for the document text extraction service."""
from __future__ import annotations

import io

import pytest

from app.services.extraction_service import extract_text, _extract_docx, _extract_pdf


def test_extract_text_routes_pdf_by_mime():
    """extract_text should handle application/pdf mime type without raising."""
    # Minimal invalid PDF — extraction will fail gracefully and return error text
    result = extract_text(b"%PDF-1.4", "application/pdf", "test.pdf")
    assert "extracted_text" in result
    assert "page_count" in result


def test_extract_text_routes_pdf_by_extension():
    """extract_text should detect PDF by .pdf extension even with wrong mime."""
    result = extract_text(b"%PDF-1.4", "application/octet-stream", "document.pdf")
    assert "extracted_text" in result


def test_extract_text_image_returns_placeholder():
    """Image files should return a placeholder string, not fail."""
    result = extract_text(b"\xff\xd8\xff", "image/jpeg", "photo.jpg")
    assert result["extracted_text"] == "[Image file: photo.jpg]"
    assert result["page_count"] == 1


def test_extract_text_unsupported_type_returns_placeholder():
    """Unknown file types should return a placeholder, not raise."""
    result = extract_text(b"random bytes", "application/x-custom", "data.xyz")
    assert "unsupported type" in result["extracted_text"]
    assert result["page_count"] == 1


def test_extract_text_extraction_failure_returns_error_text():
    """Corrupt input should return extraction error message, not raise."""
    result = extract_text(b"not a real pdf", "application/pdf", "bad.pdf")
    assert "extracted_text" in result
    # page_count should be 0 on failure
    assert result["page_count"] == 0


def test_extract_docx_returns_text():
    """_extract_docx should return paragraph text from a valid DOCX."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("Hello CMMC world")
    doc.add_paragraph("Control 3.1.1 implemented")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    result = _extract_docx(buf.getvalue())
    assert "Hello CMMC world" in result["extracted_text"]
    assert "Control 3.1.1 implemented" in result["extracted_text"]
    assert result["page_count"] == 1


def test_extract_docx_empty_doc():
    """Empty DOCX should return empty string without error."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    result = _extract_docx(buf.getvalue())
    assert result["extracted_text"] == ""
    assert result["page_count"] == 1
