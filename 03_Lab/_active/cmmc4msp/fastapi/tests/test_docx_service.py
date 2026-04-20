"""Tests for docx_service — markdown → DOCX conversion."""
from __future__ import annotations

import io
import pytest

# Import will fail until service is created — that's the RED state.
from app.services.docx_service import markdown_to_docx


def test_markdown_to_docx_returns_bytes():
    """Output is non-empty bytes."""
    result = markdown_to_docx("# Hello World", org_name="Acme Corp", control_id="3.1.1")
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_markdown_to_docx_contains_heading():
    """Output can be loaded by python-docx Document."""
    from docx import Document

    result = markdown_to_docx("# My Policy\n\nSome content here.", org_name="Test Org")
    doc = Document(io.BytesIO(result))
    # Should have at least the title paragraph
    assert len(doc.paragraphs) > 0


def test_markdown_to_docx_handles_bullet_lists():
    """Bullet list items don't crash and produce valid bytes."""
    md = "## Section\n\n- Item one\n- Item two\n- Item three\n"
    result = markdown_to_docx(md)
    assert isinstance(result, bytes)
    assert len(result) > 0

    from docx import Document
    doc = Document(io.BytesIO(result))
    assert len(doc.paragraphs) > 0


def test_markdown_to_docx_empty_markdown():
    """Empty markdown produces valid non-crashing bytes."""
    result = markdown_to_docx("", org_name="", control_id="")
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_markdown_to_docx_h1_h2_h3():
    """H1, H2, H3 heading levels all parse without crashing."""
    md = (
        "# Top Level\n"
        "## Second Level\n"
        "### Third Level\n"
        "Some body text.\n"
    )
    result = markdown_to_docx(md, org_name="ACME", control_id="3.1.2")
    from docx import Document
    doc = Document(io.BytesIO(result))
    assert len(doc.paragraphs) > 0


def test_markdown_to_docx_bold_inline():
    """Inline **bold** text is handled without crashing."""
    md = (
        "## Policy\n"
        "This is **very important** text.\n"
        "**Standalone bold line**\n"
        "Normal line with **embedded bold** in middle.\n"
    )
    result = markdown_to_docx(md, org_name="BoldCorp")
    assert isinstance(result, bytes)
    assert len(result) > 0
