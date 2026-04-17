"""Markdown to DOCX conversion using python-docx."""
from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(markdown: str, org_name: str = "", control_id: str = "") -> bytes:
    """
    Convert structured markdown to a styled DOCX document.
    Handles: # headings (H1-H3), **bold**, bullet lists (- item), plain paragraphs.
    Returns raw bytes suitable for MinIO upload.
    """
    doc = Document()

    # Title
    title = doc.add_heading(
        f"{org_name} Policy Document" if org_name else "Policy Document",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if control_id:
        sub = doc.add_paragraph(f"Control: {control_id}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer

    for line in markdown.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
        else:
            # Inline bold: split on **text** and emit mixed runs
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
